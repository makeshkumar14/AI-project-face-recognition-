"""
Generate a professional Word document report for the Face Recognition Attendance System.
Run: python generate_report.py
Output: FaceAttend_Project_Report.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'FaceAttend_Project_Report.docx')

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, header_color='2D3748', alt_color='F7FAFC'):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'
            if r_idx % 2 == 1:
                set_cell_shading(cell, alt_color)

    doc.add_paragraph()  # spacing
    return table

def build_report():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    # ==========================================================
    # TITLE PAGE
    # ==========================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI-Powered Face Recognition\nAttendance System')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Complete Project Report')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Project name
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Project: FaceAttend')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run('Technology: Python | Flask | PyTorch | OpenCV | SQLite')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.font.name = 'Calibri'

    doc.add_page_break()

    # ==========================================================
    # TABLE OF CONTENTS (manual)
    # ==========================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. Technology Stack',
        '3. System Architecture',
        '4. Features',
        '5. Database Schema',
        '6. API Endpoints',
        '7. Enrolled Students',
        '8. How to Run',
        '9. Future Enhancements',
        '10. Strengths & Limitations',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_page_break()

    # ==========================================================
    # 1. PROJECT OVERVIEW
    # ==========================================================
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        'FaceAttend is an AI-powered attendance management system that uses real-time face '
        'recognition to automatically mark student attendance. Faculty members start an attendance '
        'session, the system activates the webcam, detects and recognizes enrolled student faces, '
        'and records attendance to a database — eliminating manual roll calls entirely.'
    )
    doc.add_paragraph(
        'The system employs a dual-engine architecture: an advanced deep learning pipeline '
        '(MTCNN + FaceNet) as the primary recognizer, with an OpenCV LBPH-based fallback for '
        'environments where PyTorch is unavailable. A multi-frame voting mechanism ensures '
        'high precision by requiring consistent recognition across multiple video frames before '
        'confirming a student\'s identity.'
    )

    # ==========================================================
    # 2. TECHNOLOGY STACK
    # ==========================================================
    doc.add_heading('2. Technology Stack', level=1)

    doc.add_heading('2.1 Backend', level=2)
    add_styled_table(doc,
        ['Technology', 'Version', 'Purpose'],
        [
            ['Python', '3.9+', 'Core programming language'],
            ['Flask', '≥ 2.3.0', 'Web framework (REST API + template rendering)'],
            ['flask-cors', '≥ 4.0.0', 'Cross-Origin Resource Sharing for API endpoints'],
            ['Werkzeug', '≥ 2.3.0', 'Password hashing (generate_password_hash, check_password_hash)'],
            ['SQLite', 'Built-in', 'Lightweight relational database (no external DB server needed)'],
        ]
    )

    doc.add_heading('2.2 AI / Machine Learning', level=2)
    add_styled_table(doc,
        ['Technology', 'Purpose'],
        [
            ['MTCNN (facenet-pytorch)', 'Multi-Task Cascaded CNN for face detection & alignment'],
            ['FaceNet / InceptionResnetV1', '512-dimensional face embedding generation (pretrained on VGGFace2)'],
            ['PyTorch', 'Deep learning framework (CPU mode)'],
            ['scikit-learn', 'Cosine similarity computation for face matching'],
            ['OpenCV (opencv-contrib-python ≥ 4.8)', 'Webcam capture, image processing, Haar Cascade fallback detection'],
            ['LBPH Face Recognizer (OpenCV)', 'Fallback recognition when deep learning models unavailable'],
            ['face_recognition (dlib, optional)', 'Legacy alternate recognition library'],
        ]
    )

    doc.add_heading('2.3 Frontend', level=2)
    add_styled_table(doc,
        ['Technology', 'Purpose'],
        [
            ['HTML5 + CSS3 + JavaScript', 'Frontend UI (dashboard, login, forms)'],
            ['Jinja2 Templates', '8 server-rendered pages (role select, dashboards, login, etc.)'],
            ['Inter (Google Fonts)', 'Modern typography'],
            ['Custom CSS', 'Animated backgrounds, glassmorphism, particles, responsive layout'],
        ]
    )

    doc.add_heading('2.4 Data & Export', level=2)
    add_styled_table(doc,
        ['Technology', 'Purpose'],
        [
            ['openpyxl', 'Excel (.xlsx) export with styled formatting'],
            ['csv (stdlib)', 'CSV export for attendance records'],
            ['NumPy', 'Numerical operations on face embeddings'],
            ['Pillow', 'Image loading and format conversion'],
        ]
    )

    # ==========================================================
    # 3. SYSTEM ARCHITECTURE
    # ==========================================================
    doc.add_heading('3. System Architecture', level=1)

    doc.add_heading('3.1 Architecture Overview', level=2)
    arch_items = [
        ('Browser (HTML/CSS/JS)', 'User interacts with the frontend dashboard'),
        ('Flask Backend (app.py)', 'Routes requests, serves templates, streams video feed'),
        ('Authentication Module (routes/auth.py)', 'Faculty & student login/logout via REST API'),
        ('Attendance Module (routes/attendance.py)', 'Session control, attendance data retrieval, CSV/Excel export'),
        ('Attendance Logic (attendance_logic.py)', 'Thread-safe session management, business rules, duplicate prevention'),
        ('Advanced Face Recognition', 'MTCNN detection → FaceNet embedding → cosine similarity → multi-frame voting'),
        ('Fallback Face Recognition', 'Haar Cascade detection → LBPH recognizer (when PyTorch unavailable)'),
        ('Database (models.py → SQLite)', 'Students, faculty, and attendance CRUD operations'),
        ('Embedding Store (.npy files)', 'Pre-computed 512D face embeddings for each enrolled student'),
    ]
    for name, desc in arch_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}: ')
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(desc).font.size = Pt(10.5)

    doc.add_heading('3.2 Key Files', level=2)
    add_styled_table(doc,
        ['File', 'Lines', 'Role'],
        [
            ['app.py', '574', 'Main Flask application — routes, video feed, Excel export'],
            ['advanced_face_recognition.py', '893', 'MTCNN detector, FaceNet embedder, voting system, enrollment'],
            ['face_recognition_module.py', '475', 'Haar Cascade + LBPH fallback recognition, webcam capture'],
            ['models.py', '383', 'SQLite database schema, CRUD operations'],
            ['attendance_logic.py', '352', 'Session management, thread-safe attendance marking'],
            ['routes/auth.py', '166', 'Login/logout/session API endpoints'],
            ['routes/attendance.py', '431', 'Start/stop attendance, export CSV, student history'],
            ['enroll_students.py', '77', 'Batch enrollment script'],
        ]
    )

    # ==========================================================
    # 4. FEATURES
    # ==========================================================
    doc.add_heading('4. Features', level=1)

    doc.add_heading('4.1 Core AI Features', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Description'],
        [
            ['1', 'Real-Time Face Recognition', 'Webcam-based live face detection and identification using MTCNN + FaceNet'],
            ['2', 'Automatic Attendance Marking', 'Recognized students are automatically marked PRESENT in the database'],
            ['3', 'Multi-Frame Voting', '5-frame voting buffer — student must be recognized in ≥ 3/5 frames to confirm (reduces false positives)'],
            ['4', 'Dual Recognition Engine', 'Advanced (MTCNN + FaceNet) with fallback to OpenCV LBPH if PyTorch is unavailable'],
            ['5', 'Student Enrollment', 'Batch enrollment from image folders → face embeddings stored as .npy files'],
            ['6', 'Data Augmentation', 'Brightness, rotation, zoom, and flip augmentations during enrollment'],
        ]
    )

    doc.add_heading('4.2 User-Facing Features', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Description'],
        [
            ['7', 'Role-Based Login', 'Separate login flows for Faculty and Students'],
            ['8', 'Faculty Dashboard', 'Start/stop attendance sessions, select subject/class/section/period'],
            ['9', 'Student Dashboard', 'View personal attendance history and percentage'],
            ['10', 'Live Video Feed', 'MJPEG streaming with bounding boxes and confidence labels overlaid'],
            ['11', 'Session Management', 'Thread-safe attendance sessions with duplicate-mark prevention'],
            ['12', 'Manual Override', 'Faculty can manually mark students present/absent'],
        ]
    )

    doc.add_heading('4.3 Data & Reporting Features', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Description'],
        [
            ['13', 'Excel Export', 'Styled .xlsx reports with color-coded present/absent rows'],
            ['14', 'CSV Export', 'Lightweight CSV download with filters (subject, section, period, date)'],
            ['15', 'Attendance History API', 'Per-student attendance history with filtering'],
            ['16', 'Auto-Absent Marking', 'When session stops, all unrecognized students are auto-marked ABSENT'],
        ]
    )

    doc.add_heading('4.4 AI & Security Features', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Description'],
        [
            ['17', 'Threshold Optimization', 'Automatic calculation of optimal similarity threshold targeting 99% precision'],
            ['18', 'High-Precision Mode', 'Conservative threshold (0.55 cosine similarity) — rejects uncertain matches'],
            ['19', 'False Result Logging', 'Logs false positive/negative frames for future retraining'],
            ['20', 'Password Hashing', 'Werkzeug-based secure password hashing for faculty accounts'],
            ['21', 'Session-Based Auth', 'Flask session management with CORS protection'],
        ]
    )

    # ==========================================================
    # 5. DATABASE SCHEMA
    # ==========================================================
    doc.add_heading('5. Database Schema', level=1)
    doc.add_paragraph(
        'The system uses SQLite with three core tables. A unique constraint on the attendance table '
        'prevents duplicate marks for the same student/subject/section/period/date combination.'
    )

    doc.add_heading('5.1 Students Table', level=2)
    add_styled_table(doc,
        ['Column', 'Type', 'Constraint'],
        [
            ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'],
            ['name', 'TEXT', 'NOT NULL'],
            ['roll_no', 'TEXT', 'UNIQUE NOT NULL'],
            ['section', 'TEXT', 'NOT NULL'],
            ['image_path', 'TEXT', ''],
            ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP'],
        ]
    )

    doc.add_heading('5.2 Faculty Table', level=2)
    add_styled_table(doc,
        ['Column', 'Type', 'Constraint'],
        [
            ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'],
            ['name', 'TEXT', 'NOT NULL'],
            ['email', 'TEXT', 'UNIQUE NOT NULL'],
            ['password_hash', 'TEXT', 'NOT NULL'],
            ['department', 'TEXT', ''],
            ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP'],
        ]
    )

    doc.add_heading('5.3 Attendance Table', level=2)
    add_styled_table(doc,
        ['Column', 'Type', 'Constraint'],
        [
            ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'],
            ['student_id', 'INTEGER', 'FOREIGN KEY → students(id)'],
            ['subject', 'TEXT', 'NOT NULL'],
            ['section', 'TEXT', 'NOT NULL'],
            ['period', 'TEXT', 'NOT NULL'],
            ['date', 'TEXT', 'NOT NULL'],
            ['time', 'TEXT', 'NOT NULL'],
            ['status', 'TEXT', 'CHECK (PRESENT / ABSENT)'],
            ['confidence', 'REAL', 'DEFAULT 0.0'],
            ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP'],
        ]
    )
    p = doc.add_paragraph()
    run = p.add_run('Unique constraint: ')
    run.bold = True
    p.add_run('(student_id, subject, section, period, date) — prevents duplicate attendance entries.')

    # ==========================================================
    # 6. API ENDPOINTS
    # ==========================================================
    doc.add_heading('6. API Endpoints', level=1)
    add_styled_table(doc,
        ['Method', 'Endpoint', 'Purpose'],
        [
            ['POST', '/api/login/faculty', 'Faculty login'],
            ['POST', '/api/login/student', 'Student login'],
            ['POST', '/api/logout', 'Logout'],
            ['GET', '/api/session', 'Current session info'],
            ['POST', '/api/start_attendance', 'Start attendance session'],
            ['POST', '/api/stop_attendance', 'Stop session (marks absent)'],
            ['POST', '/api/mark_present', 'Manual mark present'],
            ['POST', '/api/mark_absent', 'Manual mark absent'],
            ['GET', '/api/get_attendance', 'Get attendance (with filters)'],
            ['GET', '/api/export_csv', 'Download CSV report'],
            ['GET', '/api/export_excel', 'Download Excel report'],
            ['GET', '/api/students', 'List all students'],
            ['GET', '/api/student_history/<roll_no>', 'Student attendance history'],
            ['POST', '/api/sync_students', 'Sync enrolled students to DB'],
            ['GET', '/video_feed', 'Live MJPEG video stream'],
        ]
    )

    # ==========================================================
    # 7. ENROLLED STUDENTS
    # ==========================================================
    doc.add_heading('7. Enrolled Students (Current Dataset)', level=1)
    doc.add_paragraph('7 students are currently enrolled with face images in the dataset:')
    add_styled_table(doc,
        ['#', 'Student Name', 'Dataset Folder'],
        [
            ['1', 'Boopathi', 'dataset/Boopathi/'],
            ['2', 'Hasvandh', 'dataset/Hasvandh/'],
            ['3', 'Kavin', 'dataset/Kavin/'],
            ['4', 'Makeshkumar', 'dataset/Makeshkumar/'],
            ['5', 'Parkavan', 'dataset/Parkavan/'],
            ['6', 'Ratchagan', 'dataset/Ratchagan/'],
            ['7', 'Rohanbala', 'dataset/Rohanbala/'],
        ]
    )

    # ==========================================================
    # 8. HOW TO RUN
    # ==========================================================
    doc.add_heading('8. How to Run', level=1)

    steps = [
        ('Step 1: Install Core Dependencies',
         'pip install -r requirements.txt'),
        ('Step 2: Install PyTorch (to temp folder for Windows)',
         'pip install torch torchvision --index-url https://download.pytorch.org/whl/cpu --target "%TEMP%\\torch_temp"\n'
         'pip install facenet-pytorch scikit-learn scipy joblib threadpoolctl --no-deps --target "%TEMP%\\torch_temp"'),
        ('Step 3: Enroll Students',
         'python enroll_students.py'),
        ('Step 4: Start the Application',
         'python app.py\n→ Server runs at http://localhost:5000'),
    ]
    for title, cmd in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        code_p = doc.add_paragraph()
        code_run = code_p.add_run(cmd)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9.5)
        code_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    # ==========================================================
    # 9. FUTURE ENHANCEMENTS
    # ==========================================================
    doc.add_heading('9. Future Enhancements & Updates', level=1)

    doc.add_heading('9.1 High Priority', level=2)
    add_styled_table(doc,
        ['#', 'Enhancement', 'Description'],
        [
            ['1', 'GPU Acceleration', 'Enable CUDA/GPU support for faster face recognition (currently CPU-only)'],
            ['2', 'Anti-Spoofing / Liveness Detection', 'Add liveness checks (blink detection, depth analysis) to prevent photo-based fraud'],
            ['3', 'Multi-Camera Support', 'Support multiple cameras for large classrooms'],
            ['4', 'Cloud Deployment', 'Deploy on AWS/GCP/Azure with a production WSGI server (Gunicorn/uWSGI)'],
            ['5', 'Proper Authentication', 'JWT token-based auth, bcrypt hashing, password reset, proper student password validation'],
        ]
    )

    doc.add_heading('9.2 Feature Additions', level=2)
    add_styled_table(doc,
        ['#', 'Enhancement', 'Description'],
        [
            ['6', 'Mobile App', 'Companion mobile app (React Native/Flutter) for student self-check and notifications'],
            ['7', 'Email/SMS Notifications', 'Automatically notify parents/students about absences'],
            ['8', 'Attendance Analytics Dashboard', 'Charts, trends, per-subject analysis, weekly attendance graphs'],
            ['9', 'Bulk Student Import', 'CSV/Excel upload for batch student registration'],
            ['10', 'QR Code Fallback', 'Secondary attendance method via QR codes when face recognition fails'],
        ]
    )

    doc.add_heading('9.3 AI Improvements', level=2)
    add_styled_table(doc,
        ['#', 'Enhancement', 'Description'],
        [
            ['11', 'Periodic Model Retraining', 'Use logged false positive/negative data to retrain and improve accuracy'],
            ['12', 'Mask Detection', 'Handle face recognition with masks (partial face matching)'],
            ['13', 'ArcFace/CosFace Embeddings', 'Upgrade from FaceNet to more modern ArcFace models for better accuracy'],
            ['14', 'Emotion & Engagement Detection', 'Detect student attention levels during class'],
            ['15', 'Edge Deployment', 'Run on Raspberry Pi or NVIDIA Jetson for low-cost, portable deployment'],
        ]
    )

    doc.add_heading('9.4 System & Infrastructure', level=2)
    add_styled_table(doc,
        ['#', 'Enhancement', 'Description'],
        [
            ['16', 'PostgreSQL/MySQL Migration', 'Replace SQLite with a production-grade database for concurrent access'],
            ['17', 'Docker Containerization', 'Dockerize the application for easy deployment and consistency'],
            ['18', 'Role-Based Access Control (RBAC)', 'Fine-grained permissions (admin, faculty, student, viewer)'],
            ['19', 'Audit Logging', 'Track all actions (session starts, manual overrides, etc.)'],
            ['20', 'Automated Testing', 'Unit tests, integration tests, and CI/CD pipeline'],
            ['21', 'WebSocket for Live Updates', 'Replace polling with WebSockets for real-time dashboard updates'],
            ['22', 'Offline Mode', 'Cache attendance data locally when internet is unavailable and sync later'],
        ]
    )

    # ==========================================================
    # 10. STRENGTHS & LIMITATIONS
    # ==========================================================
    doc.add_heading('10. Strengths & Limitations', level=1)

    doc.add_heading('Strengths', level=2)
    strengths = [
        'Dual-engine architecture — graceful fallback from deep learning to OpenCV if PyTorch is unavailable',
        'High-precision mode with multi-frame voting reduces false attendance marking',
        'Automatic threshold optimization adapts to enrolled faces automatically',
        'Data augmentation during enrollment improves recognition under varying conditions',
        'Self-contained — SQLite + embedded models, no external services required',
        'Clean modular design — separate modules for recognition, attendance, routing, and models',
    ]
    for s in strengths:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s).font.size = Pt(10.5)

    doc.add_heading('Limitations', level=2)
    limitations = [
        'CPU-only — no GPU acceleration, may be slow with many students',
        'Single-session — only one attendance session can run at a time (global singleton)',
        'No liveness detection — vulnerable to photo spoofing attacks',
        'SQLite — not suitable for concurrent multi-user production environments',
        'Flask dev server — not production-ready (no Gunicorn/Nginx)',
        'No automated tests — no unit or integration test suite',
    ]
    for l in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(l).font.size = Pt(10.5)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"\n{'='*60}")
    print(f"Report generated successfully!")
    print(f"{'='*60}")
    print(f"File: {OUTPUT_PATH}")
    print(f"{'='*60}")


if __name__ == '__main__':
    build_report()
